import fitz  # pymupdf
import re
from typing import Dict, Optional
import docx
import io
import base64
import os
from openai import OpenAI

class ResumeParser:
    def __init__(self):
        self.section_keywords = {
            "summary": ["professional summary", "summary", "profile", "objective"],
            "skills": ["skills", "technical skills", "technologies", "core competencies"],
            "experience": ["experience", "work experience", "employment history", "professional experience"],
            "education": ["education", "academic background", "qualifications"],
            "certifications": ["certifications", "certificates", "credentials"],
            "projects": ["projects", "key projects"]
        }

    def parse(self, file_path_or_bytes, filename: str = "resume.pdf") -> Dict:
        """
        Parses a PDF, DOCX, or Image file and extracts text, page count, and sections.
        """
        try:
            filename_lower = filename.lower()
            if filename_lower.endswith('.docx'):
                return self._parse_docx(file_path_or_bytes, filename)
            elif filename_lower.endswith(('.png', '.jpg', '.jpeg')):
                return self._parse_image(file_path_or_bytes, filename)
            else:
                return self._parse_pdf(file_path_or_bytes, filename)
        except Exception as e:
            print(f"Error parsing resume: {e}")
            return {
                "filename": filename,
                "page_count": 0,
                "text_content": "",
                "text": "",
                "highlighted_tokens": [],
                "sections": {}
            }

    def _normalize_whitespace(self, text: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
        return "\n".join([line for line in lines if line])

    def _parse_pdf(self, file_path_or_bytes, filename: str) -> Dict:
        if isinstance(file_path_or_bytes, str):
            doc = fitz.open(file_path_or_bytes)
        else:
            doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
        page_count = len(doc)
        full_text = ""
        highlighted_tokens = set()
        
        for page in doc:
            full_text += page.get_text() + "\n"
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if b.get('type') == 0:
                    for line in b.get("lines", []):
                        for span in line.get("spans", []):
                            text_span = span.get("text", "").strip()
                            if not text_span: continue
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 16)
                            is_caps = text_span.isupper() and len(text_span) > 3
                            if is_bold or is_caps:
                                clean_tok = re.sub(r'[^\w\s]', '', text_span).lower()
                                highlighted_tokens.add(clean_tok)

        full_text = self._normalize_whitespace(full_text)
        sections = self._extract_sections(full_text)
        flat_text = re.sub(r"\s+", " ", full_text).strip()
        
        return {
            "filename": filename,
            "file_path": "memory",
            "page_count": page_count,
            "text_content": full_text,
            "text": flat_text,
            "highlighted_tokens": list(highlighted_tokens),
            "sections": sections
        }

    def _parse_docx(self, file_path_or_bytes, filename: str) -> Dict:
        if isinstance(file_path_or_bytes, str):
            doc = docx.Document(file_path_or_bytes)
        else:
            doc = docx.Document(io.BytesIO(file_path_or_bytes))
            
        full_text = "\n".join([para.text for para in doc.paragraphs])
        highlighted_tokens = set()
        
        for para in doc.paragraphs:
            for run in para.runs:
                text_span = run.text.strip()
                if not text_span: continue
                is_bold = run.bold
                is_caps = text_span.isupper() and len(text_span) > 3
                if is_bold or is_caps:
                    clean_tok = re.sub(r'[^\w\s]', '', text_span).lower()
                    highlighted_tokens.add(clean_tok)
                    
        full_text = self._normalize_whitespace(full_text)
        sections = self._extract_sections(full_text)
        flat_text = re.sub(r"\s+", " ", full_text).strip()
        
        return {
            "filename": filename,
            "file_path": "memory",
            "page_count": 1, # DOCX doesn't have easy page count in this lib
            "text_content": full_text,
            "text": flat_text,
            "highlighted_tokens": list(highlighted_tokens),
            "sections": sections
        }

    def _parse_image(self, file_path_or_bytes, filename: str) -> Dict:
        if isinstance(file_path_or_bytes, str):
            with open(file_path_or_bytes, "rb") as image_file:
                base_64_image = base64.b64encode(image_file.read()).decode('utf-8')
        else:
            base_64_image = base64.b64encode(file_path_or_bytes).decode('utf-8')
            
        ext = filename.split(".")[-1].lower()
        if ext == 'jpg': ext = 'jpeg'
        mime_type = f"image/{ext}"

        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        prompt = '''
        You are a resume extractor. Read this resume image and extract its entire text content exactly as it appears. 
        Formatting should be preserved where possible.
        '''
        
        response = client.chat.completions.create(
            model="gpt-5",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base_64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=3000
        )
        
        full_text = response.choices[0].message.content or ""
        full_text = self._normalize_whitespace(full_text)
        sections = self._extract_sections(full_text)
        flat_text = re.sub(r"\s+", " ", full_text).strip()
        
        return {
            "filename": filename,
            "file_path": "memory",
            "page_count": 1,
            "text_content": full_text,
            "text": flat_text,
            "highlighted_tokens": [], # Hard to extract easily from basic vision without advanced prompting
            "sections": sections
        }

    def _extract_sections(self, text: str) -> Dict[str, str]:
        lines = text.splitlines()
        if not lines:
            return {}

        section_starts = []
        for index, line in enumerate(lines):
            normalized = re.sub(r"[^a-z0-9 ]", "", line.lower()).strip()
            for section, keywords in self.section_keywords.items():
                if any(normalized == keyword or normalized.startswith(keyword + " ") for keyword in keywords):
                    section_starts.append((section, index))
                    break

        if not section_starts:
            return {}

        extracted = {}
        for position, (section_name, start_index) in enumerate(section_starts):
            end_index = section_starts[position + 1][1] if position + 1 < len(section_starts) else len(lines)
            content = "\n".join(lines[start_index:end_index]).strip()
            extracted[section_name] = content

        return extracted
